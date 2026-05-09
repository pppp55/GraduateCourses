from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from embedding_utils import spearman_correlation


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS = PROJECT_ROOT / "results"
REPORT_DIR = PROJECT_ROOT / "report"
STUDENT_ID = "225040065"


def load_results() -> dict:
    data = {}
    for model in ["cbow", "skipgram"]:
        knn_summary = pd.read_csv(RESULTS / f"{model}_knn_summary.csv")
        simlex = pd.read_csv(RESULTS / f"{model}_simlex_all.csv")
        analogy = pd.read_csv(RESULTS / f"{model}_analogy_all.csv")
        analogy_category = pd.read_csv(RESULTS / f"{model}_analogy_by_category.csv")

        simlex_cov = simlex[simlex["covered"] == 1].copy()
        analogy_cov = analogy[analogy["covered"] == 1].copy()
        analogy_cov["correct"] = pd.to_numeric(analogy_cov["correct"], errors="coerce")

        data[model] = {
            "knn_summary": knn_summary,
            "knn_detail": pd.read_csv(RESULTS / f"{model}_knn_detail_words.csv"),
            "simlex": simlex,
            "simlex_sample": pd.read_csv(RESULTS / f"{model}_simlex_sample_20.csv"),
            "analogy": analogy,
            "analogy_sample": pd.read_csv(RESULTS / f"{model}_analogy_sample_10.csv"),
            "analogy_category": analogy_category,
            "knn_avg": pd.to_numeric(knn_summary["average_similarity"], errors="coerce").mean(),
            "knn_covered": int(knn_summary["covered"].sum()),
            "simlex_covered": len(simlex_cov),
            "simlex_spearman": spearman_correlation(
                simlex_cov["standard_similarity"].astype(float).tolist(),
                simlex_cov["scaled_similarity"].astype(float).tolist(),
            ),
            "analogy_covered": len(analogy_cov),
            "analogy_correct": int(analogy_cov["correct"].sum()),
            "analogy_accuracy": float(analogy_cov["correct"].mean()),
        }
    return data


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, size=8)
        set_cell_shading(hdr[i], "EAF2F8")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str = ""):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def neighbor_list(df: pd.DataFrame, query: str) -> str:
    sub = df[df["query"] == query].sort_values("rank")
    return "; ".join(f"{r.neighbor} ({float(r.cosine_similarity):.3f})" for r in sub.itertuples())


def build_markdown(data: dict) -> str:
    cbow = data["cbow"]
    skip = data["skipgram"]
    detail_words = sorted(cbow["knn_detail"]["query"].unique())

    lines = [
        "# CSC5010 Artificial Intelligence Project 1 and 2 Report",
        "",
        f"Student ID: {STUDENT_ID}",
        "",
        "## 1. Overview",
        "",
        "This project trains and evaluates two Word2Vec models, Continuous Bag of Words (CBOW) and Skip-Gram, on the NLTK Reuters corpus. Both models use 64-dimensional embeddings, context window size 2, batch size 1024, Adam optimizer with learning rate 0.001, and 10 training epochs. The final embedding files are saved as `embeddings/cbow.vec` and `embeddings/skipgram.vec`.",
        "",
        "The evaluation contains three parts: K-nearest neighbor word similarity, SimLex-999 golden standard correlation, and Mikolov-style analogy reasoning. Random report examples are selected using student ID 225040065 as the reproducible seed.",
        "",
        "## 2. Training Results",
        "",
        "| Model | Vocabulary | Dimension | Final loss | Embedding file |",
        "|---|---:|---:|---:|---|",
        "| CBOW | 31081 | 64 | 6462.34 | embeddings/cbow.vec |",
        "| Skip-Gram | 31081 | 64 | 37952.81 | embeddings/skipgram.vec |",
        "",
        "The CBOW loss is lower in absolute value because each training example predicts one center word from its context, while Skip-Gram produces many more center-context pairs. The two loss values therefore should not be compared directly as model quality scores. Evaluation results below provide the more meaningful comparison.",
        "",
        "## 3. K-Nearest Neighbor Evaluation",
        "",
        f"Both models covered 49 of the 50 required query words. The missing query is caused by vocabulary coverage in the Reuters-trained embeddings. CBOW reached an overall average top-10 cosine similarity of {cbow['knn_avg']:.6f}; Skip-Gram reached {skip['knn_avg']:.6f}. The four seed-selected detailed query words are {', '.join(detail_words)}.",
        "",
        "| Query | CBOW top-10 neighbors | Skip-Gram top-10 neighbors |",
        "|---|---|---|",
    ]
    for word in detail_words:
        lines.append(f"| {word} | {neighbor_list(cbow['knn_detail'], word)} | {neighbor_list(skip['knn_detail'], word)} |")

    lines.extend([
        "",
        "The KNN results show that the models learn a mixture of semantic and topical relatedness. Query words such as `company` and `student` often retrieve words that are plausible in Reuters-style contexts rather than strict dictionary synonyms. This behavior is expected because Word2Vec learns from distributional context, and Reuters text is finance/news oriented. Skip-Gram has a slightly higher average top-10 similarity, but the qualitative difference is small.",
        "",
        "## 4. SimLex-999 Golden Standard Evaluation",
        "",
        f"Out of 999 SimLex-999 word pairs, 680 pairs are covered by both embedding vocabularies. CBOW obtains Spearman correlation {cbow['simlex_spearman']:.6f}; Skip-Gram obtains {skip['simlex_spearman']:.6f}. The low correlations indicate that cosine similarity from this relatively small Reuters training corpus does not align strongly with human lexical similarity judgments.",
        "",
        "| Model | Covered pairs | Spearman correlation |",
        "|---|---:|---:|",
        f"| CBOW | {cbow['simlex_covered']} / 999 | {cbow['simlex_spearman']:.6f} |",
        f"| Skip-Gram | {skip['simlex_covered']} / 999 | {skip['simlex_spearman']:.6f} |",
        "",
        "The sampled pairs show the main weakness: many unrelated or weakly related pairs are mapped near the middle of the scaled [0, 10] range. For example, CBOW gives `water-salt` 6.890 even though the standard score is 1.30, and Skip-Gram gives `boat-car` 6.238 while the standard score is 2.37. Positive examples also appear: both models assign relatively high scores to `essential-necessary`, `get-buy`, and `attend-arrive`. Overall, Skip-Gram is slightly better by Spearman correlation, but the margin is very small.",
        "",
        "## 5. Analogical Reasoning Evaluation",
        "",
        f"For analogy reasoning, only questions where all four words are in the vocabulary are evaluated. Both models cover 8591 questions. CBOW answers {cbow['analogy_correct']} correctly, accuracy {cbow['analogy_accuracy']:.6f}; Skip-Gram answers {skip['analogy_correct']} correctly, accuracy {skip['analogy_accuracy']:.6f}.",
        "",
        "| Model | Covered questions | Correct | Accuracy |",
        "|---|---:|---:|---:|",
        f"| CBOW | {cbow['analogy_covered']} | {cbow['analogy_correct']} | {cbow['analogy_accuracy']:.6f} |",
        f"| Skip-Gram | {skip['analogy_covered']} | {skip['analogy_correct']} | {skip['analogy_accuracy']:.6f} |",
        "",
        "The analogy task is the hardest part of the evaluation. The predicted nearest words in the sampled examples are mostly incorrect and often look like Reuters-specific tokens rather than the expected analogy answers. This suggests that the learned vector offsets are not stable enough for exact analogy completion, especially for capitals, currencies, and morphology categories. Skip-Gram performs slightly better overall, but both accuracies are low.",
        "",
        "## 6. Conclusion",
        "",
        "Skip-Gram performs slightly better than CBOW in all three aggregate metrics in this run: KNN average similarity, SimLex Spearman correlation, and analogy accuracy. However, the improvement is modest for KNN and SimLex, and both models struggle with exact analogy reasoning. The main limitations are the limited domain and scale of the Reuters corpus, the small 64-dimensional embedding size, and plain Word2Vec training without negative sampling or larger external corpora. The results still show that the models learn useful local distributional similarity, but not robust human-level semantic similarity or analogy structure.",
    ])
    return "\n".join(lines)


def build_docx(data: dict, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CSC5010 Artificial Intelligence Project 1 and 2 Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    meta = doc.add_paragraph(f"Student ID: {STUDENT_ID} | Models: CBOW and Skip-Gram | Dataset: NLTK Reuters")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cbow = data["cbow"]
    skip = data["skipgram"]
    detail_words = sorted(cbow["knn_detail"]["query"].unique())

    add_heading(doc, "1. Overview")
    add_paragraph(doc, "This project trains and evaluates two Word2Vec models: Continuous Bag of Words (CBOW) and Skip-Gram. Both models are trained on the NLTK Reuters corpus with 64-dimensional word embeddings, context window size 2, batch size 1024, Adam optimizer with learning rate 0.001, and 10 epochs.")
    add_paragraph(doc, "The three evaluation tasks are K-nearest neighbor word similarity, SimLex-999 golden standard correlation, and Mikolov-style analogy reasoning. All randomly selected report examples use student ID 225040065 as the reproducible seed.")

    add_heading(doc, "2. Training Results")
    add_table(doc, ["Model", "Vocabulary", "Dimension", "Final loss", "Embedding file"], [
        ["CBOW", "31081", "64", "6462.34", "embeddings/cbow.vec"],
        ["Skip-Gram", "31081", "64", "37952.81", "embeddings/skipgram.vec"],
    ], [1.0, 1.0, 0.9, 1.0, 2.5])
    add_paragraph(doc, "The loss values are not directly comparable because CBOW predicts one center word from surrounding context, while Skip-Gram creates more center-context training pairs. The evaluation metrics below are therefore used to compare embedding quality.")

    add_heading(doc, "3. K-Nearest Neighbor Evaluation")
    add_paragraph(doc, f"Both models cover 49 of the 50 required query words. CBOW obtains an overall average top-10 cosine similarity of {cbow['knn_avg']:.6f}, while Skip-Gram obtains {skip['knn_avg']:.6f}. The four seed-selected detailed words are {', '.join(detail_words)}.")
    rows = []
    for word in detail_words:
        rows.append([word, neighbor_list(cbow["knn_detail"], word), neighbor_list(skip["knn_detail"], word)])
    add_table(doc, ["Query", "CBOW top-10 neighbors", "Skip-Gram top-10 neighbors"], rows, [0.8, 3.0, 3.0])
    add_paragraph(doc, "The retrieved neighbors show topical similarity more clearly than strict synonymy. This is especially visible for words such as company and student, where Reuters-style context strongly affects which words appear nearby. Skip-Gram is slightly higher in average similarity, but the qualitative difference is small.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "4. SimLex-999 Golden Standard Evaluation")
    add_paragraph(doc, f"SimLex-999 contains 999 human-labeled word pairs. Both models cover 680 pairs. Cosine similarities are scaled to [0, 10] using (cosine + 1) * 5, then compared with human scores using Spearman correlation.")
    add_table(doc, ["Model", "Covered pairs", "Spearman correlation"], [
        ["CBOW", f"{cbow['simlex_covered']} / 999", f"{cbow['simlex_spearman']:.6f}"],
        ["Skip-Gram", f"{skip['simlex_covered']} / 999", f"{skip['simlex_spearman']:.6f}"],
    ], [1.2, 1.4, 1.6])
    add_paragraph(doc, "The correlations are low, which means the Reuters-trained embeddings only weakly match human lexical similarity judgments. Many word pairs are placed near the middle of the scaled range, even when their human score is low.")
    sim_rows = []
    cbow_sample = cbow["simlex_sample"].head(20)
    skip_sample = skip["simlex_sample"].head(20)
    for c, s in zip(cbow_sample.itertuples(), skip_sample.itertuples()):
        pair = f"{c.word1}-{c.word2}"
        sim_rows.append([
            pair,
            f"{float(c.standard_similarity):.2f}",
            f"{float(c.scaled_similarity):.2f}",
            f"{float(s.scaled_similarity):.2f}",
        ])
    add_table(doc, ["Pair", "Standard", "CBOW scaled", "Skip-Gram scaled"], sim_rows, [1.8, 1.0, 1.2, 1.2])
    add_paragraph(doc, "The sample includes reasonable high-similarity cases such as essential-necessary and get-buy, but also clear errors such as water-salt and boat-car receiving mid-range model scores. Skip-Gram has the better Spearman score, but only by a small margin.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "5. Analogical Reasoning Evaluation")
    add_paragraph(doc, "For each analogy A:B :: C:D, the predicted vector is embedding(B) - embedding(A) + embedding(C), and the nearest vocabulary word is searched by cosine similarity while excluding A, B, and C.")
    add_table(doc, ["Model", "Covered questions", "Correct", "Accuracy"], [
        ["CBOW", str(cbow["analogy_covered"]), str(cbow["analogy_correct"]), f"{cbow['analogy_accuracy']:.6f}"],
        ["Skip-Gram", str(skip["analogy_covered"]), str(skip["analogy_correct"]), f"{skip['analogy_accuracy']:.6f}"],
    ], [1.2, 1.6, 1.0, 1.0])

    cat_rows = []
    for _, row in cbow["analogy_category"].iterrows():
        srow = skip["analogy_category"][skip["analogy_category"]["category"] == row["category"]].iloc[0]
        cat_rows.append([
            row["category"],
            f"{int(row['correct'])}/{int(row['covered_questions'])}",
            f"{float(row['accuracy']):.4f}" if pd.notna(row["accuracy"]) else "",
            f"{int(srow['correct'])}/{int(srow['covered_questions'])}",
            f"{float(srow['accuracy']):.4f}" if pd.notna(srow["accuracy"]) else "",
        ])
    add_table(doc, ["Category", "CBOW correct", "CBOW acc.", "Skip correct", "Skip acc."], cat_rows, [2.2, 1.1, 0.9, 1.1, 0.9])
    add_paragraph(doc, "The analogy task is difficult for both models. Skip-Gram answers 31 questions correctly and CBOW answers 28. Most sampled predictions are unrelated to the expected answer, suggesting that the vector offsets learned from Reuters are not stable enough for exact analogy completion.")

    add_heading(doc, "Sample Analogy Predictions", 2)
    ana_rows = []
    for c, s in zip(cbow["analogy_sample"].head(10).itertuples(), skip["analogy_sample"].head(10).itertuples()):
        query = f"{c.word_a}:{c.word_b} :: {c.word_c}:{c.expected}"
        ana_rows.append([c.category, query, c.predicted, s.predicted])
    add_table(doc, ["Category", "Analogy", "CBOW prediction", "Skip-Gram prediction"], ana_rows, [1.7, 2.6, 1.3, 1.3])

    add_heading(doc, "Vector Plot Examples", 2)
    plot_paths = [
        RESULTS / "cbow_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png",
        RESULTS / "skipgram_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png",
    ]
    for path in plot_paths:
        if path.exists():
            doc.add_picture(str(path), width=Inches(5.8))
            cap = doc.add_paragraph(path.parent.name.replace("_", " ") + ": " + path.stem)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "The full vector plot PNGs and CSV values for five categories per model are included in the results folders.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. Conclusion")
    add_paragraph(doc, "Skip-Gram performs slightly better than CBOW in all three aggregate metrics in this run: KNN average similarity, SimLex Spearman correlation, and analogy accuracy. However, the gap is modest for KNN and SimLex, and both models struggle with exact analogy reasoning.")
    add_paragraph(doc, "The main limitations are the Reuters corpus size and domain, the 64-dimensional embedding size, and the plain Word2Vec objective. Reuters text is dominated by financial and news contexts, so many nearest neighbors reflect topical co-occurrence rather than broad semantic similarity. A larger and more balanced corpus, more epochs, tuned hyperparameters, or negative sampling would likely improve the evaluations.")
    add_paragraph(doc, "Overall, the models demonstrate the basic distributional learning behavior expected from Word2Vec, but the golden standard and analogy tasks show that this training setup is not sufficient for robust human-level semantic similarity or analogy structure.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    data = load_results()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown(data)
    (REPORT_DIR / "project1_2_report.md").write_text(markdown, encoding="utf-8")
    build_docx(data, REPORT_DIR / "project1_2_report.docx")
    print(REPORT_DIR / "project1_2_report.docx")
    print(REPORT_DIR / "project1_2_report.md")


if __name__ == "__main__":
    main()
