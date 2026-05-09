from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_report import (
    PROJECT_ROOT,
    REPORT_DIR,
    RESULTS,
    STUDENT_ID,
    add_heading,
    add_paragraph,
    add_table,
    load_results,
    neighbor_list,
)


ASSETS = PROJECT_ROOT / "report" / "assets"


def md_image(path: Path, label: str) -> str:
    rel = path.relative_to(PROJECT_ROOT).as_posix()
    return f"![{label}](../{rel})"


def sample_simlex_comment(word1: str, word2: str, standard: float, cbow_scaled: float, skip_scaled: float) -> str:
    if standard >= 6.0 and max(cbow_scaled, skip_scaled) >= 6.0:
        return "model scores move in the expected high-similarity direction"
    if standard <= 2.0 and min(cbow_scaled, skip_scaled) >= 5.0:
        return "models overestimate the relation, likely because topical context is being learned"
    if abs(cbow_scaled - skip_scaled) < 0.4:
        return "both models behave similarly on this pair"
    return "the two architectures disagree moderately on this pair"


def build_markdown(data: dict) -> str:
    cbow = data["cbow"]
    skip = data["skipgram"]
    detail_words = sorted(cbow["knn_detail"]["query"].unique())

    lines: list[str] = [
        "# CSC5010 Artificial Intelligence Project 1 and 2 Report",
        "",
        f"Student ID: {STUDENT_ID}",
        "",
        "## Abstract",
        "",
        "This report presents the training and evaluation of two Word2Vec models, Continuous Bag of Words (CBOW) and Skip-Gram, for CSC5010 Project 1 and 2. Both models are trained on the NLTK Reuters corpus and evaluated using three tasks: K-nearest neighbor similarity, SimLex-999 golden standard correlation, and analogical reasoning. The experiment is intentionally small and reproducible: the embedding dimension is 64, the context window is 2, the batch size is 1024, the optimizer is Adam with learning rate 0.001, and both models are trained for 10 epochs.",
        "",
        "The main finding is that Skip-Gram is slightly stronger than CBOW on all three aggregate metrics in this run, but the difference is small. Both models learn useful local distributional similarity, while the SimLex and analogy scores reveal clear limitations caused by the small Reuters training corpus, narrow news/finance domain, and plain Word2Vec objective.",
        "",
        "## 1. Task Definition and Dataset",
        "",
        "The goal of this project is to train two word embedding models and evaluate whether the resulting word vectors capture semantic similarity and word relation structure. The two models answer complementary self-supervised questions. CBOW predicts a center word from surrounding context words, while Skip-Gram predicts surrounding context words from a center word. In both cases, no manually labeled training data is used for learning the vectors.",
        "",
        "The training corpus is the Reuters corpus from NLTK. After preprocessing and lowercasing, the corpus contains 54,716 sentences and produces a vocabulary of 31,081 tokens. This corpus is appropriate for a compact experiment, but it is not a broad general-language corpus. Many documents are about markets, trade, companies, commodities, policy, and financial events. This domain bias matters in the evaluation because Word2Vec learns words that occur in similar contexts, not necessarily words that humans judge as synonyms.",
        "",
        "The evaluation files are the assignment-provided 50 KNN query words, `simlex-999.txt`, and `analogical reasoning task.txt`. Random examples in this report are selected with my student ID, 225040065, so that the KNN detailed words, SimLex sample pairs, and analogy examples can be reproduced by running the scripts again with the same seed.",
        "",
        "## 2. Model Training",
        "",
        "| Model | Training objective | Vocabulary | Dimension | Context window | Epochs | Final loss | Output file |",
        "|---|---|---:|---:|---:|---:|---:|---|",
        "| CBOW | context -> center word | 31,081 | 64 | 2 | 10 | 6462.34 | `embeddings/cbow.vec` |",
        "| Skip-Gram | center word -> context | 31,081 | 64 | 2 | 10 | 37952.81 | `embeddings/skipgram.vec` |",
        "",
        "The two final loss values should not be directly compared as if they were the same measurement. CBOW creates one prediction target for each center position, while Skip-Gram creates multiple center-context pairs and therefore accumulates a larger training loss. Instead, the downstream evaluation tasks provide the meaningful comparison.",
        "",
        "The trained vectors are saved in text `.vec` format. The first line records vocabulary size and embedding dimension, followed by one token and its vector per line. This format is used consistently by all evaluation scripts.",
        "",
        "## 3. Evaluation Methodology",
        "",
        "For KNN evaluation, cosine similarity is computed between each query word vector and every vocabulary vector. The query word itself is excluded, and the top 10 nearest neighbors are saved. Each query also receives an average top-10 cosine similarity, and the final KNN score is the average across all covered queries.",
        "",
        "For the golden standard evaluation, each covered SimLex-999 pair is scored using cosine similarity. Because SimLex labels are in [0, 10] while cosine similarity is in [-1, 1], the model score is normalized by `scaled = (cosine + 1) * 5`. Spearman correlation is then computed between human standard scores and scaled embedding scores.",
        "",
        "For analogical reasoning, each question has the form A:B :: C:D. The predicted vector is `embedding(B) - embedding(A) + embedding(C)`. The nearest vocabulary word to the predicted vector is searched by cosine similarity, excluding A, B, and C. A prediction is correct only if the nearest word exactly matches D.",
        "",
        "## 4. K-Nearest Neighbor Evaluation",
        "",
        f"Both models cover 49 of the 50 required query words. CBOW obtains an overall average top-10 cosine similarity of {cbow['knn_avg']:.6f}. Skip-Gram obtains {skip['knn_avg']:.6f}. The four student-ID-selected detailed words are **{', '.join(detail_words)}**.",
        "",
        md_image(ASSETS / "knn_overall.png", "KNN overall average similarity"),
        "",
        md_image(ASSETS / "knn_detail_words.png", "KNN detailed query word averages"),
        "",
        "| Query | CBOW top-10 neighbors | Skip-Gram top-10 neighbors |",
        "|---|---|---|",
    ]

    for word in detail_words:
        lines.append(f"| {word} | {neighbor_list(cbow['knn_detail'], word)} | {neighbor_list(skip['knn_detail'], word)} |")

    lines.extend([
        "",
        "The KNN results show that both models learn distributional relatedness, but this relatedness is not always the same as human synonymy. For `company`, Skip-Gram retrieves `companies`, which is a strong morphological and semantic match, while CBOW retrieves `firm` but also several news-specific or noisy tokens. This suggests that Skip-Gram can preserve some local lexical regularities, but both models are affected by rare or domain-specific Reuters vocabulary.",
        "",
        "For `help`, both models retrieve words related to support or intervention, such as `assistance`, `ease`, `solve`, and `maintain`. This is one of the more interpretable KNN cases. For `play`, the neighbors are more mixed. CBOW retrieves verbs such as `find`, `develop`, and `recommend`, while Skip-Gram retrieves words such as `deter` and `cope`; these are not synonyms, but they can occur in similar syntactic positions. For `student`, the neighbors are weak, which is reasonable because Reuters is not education-centered, so the model sees fewer clean contexts for this word.",
        "",
        "Overall, the KNN task gives the most favorable view of the embeddings because it measures local neighborhood coherence. Even when the neighbors are not exact synonyms, many of them share part of speech, topic, or usage context. Skip-Gram is slightly higher numerically, but the practical difference is modest.",
        "",
        "## 5. SimLex-999 Golden Standard Evaluation",
        "",
        f"Out of 999 SimLex-999 word pairs, 680 pairs are covered by the learned vocabulary. CBOW obtains Spearman correlation {cbow['simlex_spearman']:.6f}; Skip-Gram obtains {skip['simlex_spearman']:.6f}. These correlations are low, which means the ranking induced by the embedding similarities does not align strongly with human semantic similarity judgments.",
        "",
        "| Model | Covered pairs | Spearman correlation | Interpretation |",
        "|---|---:|---:|---|",
        f"| CBOW | {cbow['simlex_covered']} / 999 | {cbow['simlex_spearman']:.6f} | weak positive alignment |",
        f"| Skip-Gram | {skip['simlex_covered']} / 999 | {skip['simlex_spearman']:.6f} | slightly better weak alignment |",
        "",
        md_image(ASSETS / "simlex_sample_scatter.png", "SimLex sampled pairs scatter"),
        "",
        "| Pair | Standard | CBOW scaled | Skip-Gram scaled | Comment |",
        "|---|---:|---:|---:|---|",
    ])

    cbow_sample = cbow["simlex_sample"].head(20)
    skip_sample = skip["simlex_sample"].head(20)
    for c, s in zip(cbow_sample.itertuples(), skip_sample.itertuples()):
        pair = f"{c.word1}-{c.word2}"
        comment = sample_simlex_comment(c.word1, c.word2, float(c.standard_similarity), float(c.scaled_similarity), float(s.scaled_similarity))
        lines.append(f"| {pair} | {float(c.standard_similarity):.2f} | {float(c.scaled_similarity):.2f} | {float(s.scaled_similarity):.2f} | {comment} |")

    lines.extend([
        "",
        "The scatter plot makes the weakness visible: many points lie around the middle of the model-score range instead of following the diagonal. For example, `water-salt` has a low human score of 1.30 but receives a CBOW scaled score of 6.89. This kind of error is understandable in a Reuters corpus, where commodities and physical objects may appear in related news contexts even when they are not semantically similar in the SimLex sense.",
        "",
        "There are also pairs where the embedding model behaves more reasonably. Both models assign higher scores to `essential-necessary`, `get-buy`, and `attend-arrive`. These cases show that the embeddings do capture some semantic regularity, but the signal is not consistent enough to produce a high global rank correlation. Skip-Gram is slightly better, but the difference is too small to claim a large advantage.",
        "",
        "## 6. Analogical Reasoning Evaluation",
        "",
        f"The analogy dataset is much harder than KNN or SimLex because it requires vector offsets to represent stable relationships. Both models cover 8,591 analogy questions where all four words are present in the vocabulary. CBOW answers {cbow['analogy_correct']} questions correctly, and Skip-Gram answers {skip['analogy_correct']} correctly.",
        "",
        "| Model | Covered questions | Correct | Accuracy |",
        "|---|---:|---:|---:|",
        f"| CBOW | {cbow['analogy_covered']} | {cbow['analogy_correct']} | {cbow['analogy_accuracy']:.6f} |",
        f"| Skip-Gram | {skip['analogy_covered']} | {skip['analogy_correct']} | {skip['analogy_accuracy']:.6f} |",
        "",
        md_image(ASSETS / "analogy_category_accuracy.png", "Analogy category accuracy"),
        "",
        "| Category | CBOW correct / covered | CBOW acc. | Skip-Gram correct / covered | Skip-Gram acc. |",
        "|---|---:|---:|---:|---:|",
    ])

    for _, row in cbow["analogy_category"].iterrows():
        srow = skip["analogy_category"][skip["analogy_category"]["category"] == row["category"]].iloc[0]
        lines.append(
            f"| {row['category']} | {int(row['correct'])}/{int(row['covered_questions'])} | {float(row['accuracy']):.4f} | "
            f"{int(srow['correct'])}/{int(srow['covered_questions'])} | {float(srow['accuracy']):.4f} |"
        )

    lines.extend([
        "",
        "The category table shows that most analogy categories have near-zero accuracy. CBOW performs best on `gram6-nationality-adjective`, while Skip-Gram spreads its few correct answers across more categories. However, the absolute numbers are still very low. This result suggests that the learned embedding space is not organized enough for exact relation arithmetic.",
        "",
        "| Category | Analogy | Expected | CBOW predicted | Skip-Gram predicted |",
        "|---|---|---|---|---|",
    ])
    for c, s in zip(cbow["analogy_sample"].head(10).itertuples(), skip["analogy_sample"].head(10).itertuples()):
        query = f"{c.word_a}:{c.word_b} :: {c.word_c}:?"
        lines.append(f"| {c.category} | {query} | {c.expected} | {c.predicted} | {s.predicted} |")

    lines.extend([
        "",
        "Most sampled analogy predictions are not close to the expected answers. For instance, geographic and grammatical questions often return unrelated Reuters vocabulary. This is a stronger failure than the SimLex result because analogy reasoning depends on the geometry of differences between vectors, not only local similarity. A small domain-specific corpus can learn that words are nearby, but it may not learn that many word pairs share the same offset direction.",
        "",
        "### Vector Plot Examples",
        "",
        md_image(RESULTS / "cbow_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png", "CBOW analogy vector plot"),
        "",
        md_image(RESULTS / "skipgram_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png", "Skip-Gram analogy vector plot"),
        "",
        md_image(RESULTS / "cbow_analogy_plots" / "family_brother_sister_brothers.png", "CBOW family vector plot"),
        "",
        md_image(RESULTS / "skipgram_analogy_plots" / "family_brother_sister_brothers.png", "Skip-Gram family vector plot"),
        "",
        "The vector plots compare the predicted vector with the nearest searched word vector across the 64 embedding dimensions. Even when the line shapes look partially aligned, the nearest word can still be semantically wrong. This is because the nearest-neighbor search is sensitive to the entire high-dimensional geometry, and a visually similar 64-dimensional curve does not guarantee that the expected analogy answer is the nearest vocabulary item.",
        "",
        "## 7. Overall Comparison and Discussion",
        "",
        "Across the three tasks, Skip-Gram is consistently but only slightly stronger. Its KNN average is 0.499827 compared with CBOW's 0.490364. Its SimLex Spearman correlation is 0.058087 compared with CBOW's 0.055747. Its analogy accuracy is 0.003608 compared with CBOW's 0.003259. The consistency is meaningful, but the margin is not large.",
        "",
        "The results match the expected behavior of Word2Vec on a compact corpus. KNN evaluation is relatively forgiving because it asks whether nearby words are contextually related. SimLex is stricter because it compares model similarities against human lexical similarity. Analogy reasoning is strictest because it requires relational directions such as country-capital, adjective-adverb, comparative, superlative, and plural transformations to be linearly encoded in the vector space.",
        "",
        "The main limitation is not that the implementation failed; the models trained correctly and produced valid embeddings. The limitation is that a 64-dimensional Word2Vec model trained for 10 epochs on Reuters does not have enough broad linguistic evidence to solve the hardest semantic tasks well. A larger corpus such as Wikipedia or Common Crawl, more training time, better handling of rare tokens, negative sampling, and hyperparameter tuning would likely improve the golden standard and analogy results.",
        "",
        "## 8. Conclusion",
        "",
        "This project successfully trained and evaluated CBOW and Skip-Gram embeddings. The KNN results show that both models learn meaningful local distributional neighborhoods. The SimLex results show weak alignment with human semantic similarity, and the analogy results show that the learned vector space does not reliably preserve exact relational offsets. Skip-Gram is the better model in this run, but the improvement over CBOW is small.",
        "",
        "In conclusion, the experiment demonstrates both the strength and the limitation of simple Word2Vec. It is effective at learning contextual similarity from unlabeled text, but robust semantic similarity and analogy reasoning require broader data, stronger objectives, and more careful tuning than this compact assignment setup provides.",
    ])
    return "\n".join(lines)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.3)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)
    return doc


def add_picture(doc: Document, path: Path, caption: str, width: float = 5.8) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True


def build_docx(data: dict, output: Path) -> None:
    doc = setup_doc()
    cbow = data["cbow"]
    skip = data["skipgram"]
    detail_words = sorted(cbow["knn_detail"]["query"].unique())

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CSC5010 Artificial Intelligence Project 1 and 2 Report")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(31, 78, 121)
    meta = doc.add_paragraph(f"Student ID: {STUDENT_ID} | CBOW and Skip-Gram Word2Vec | NLTK Reuters")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Abstract")
    add_paragraph(doc, "This report presents the training and evaluation of CBOW and Skip-Gram Word2Vec embeddings on the NLTK Reuters corpus. The models are evaluated with K-nearest neighbor similarity, SimLex-999 golden standard correlation, and Mikolov-style analogy reasoning. Skip-Gram is slightly stronger across the aggregate metrics, but both models show clear limitations on human similarity ranking and exact analogy completion.")
    add_heading(doc, "1. Task Definition and Dataset")
    add_paragraph(doc, "The project objective is to train two neural word embedding models and evaluate whether the learned vectors encode semantic similarity and word relations. CBOW predicts a center word from its context, while Skip-Gram predicts context words from a center word. Both methods learn from raw text without manual labels.")
    add_paragraph(doc, "The training corpus is NLTK Reuters. It contains 54,716 sentences after loading and lowercasing, producing a vocabulary of 31,081 tokens. Reuters is useful for a compact assignment, but it is heavily biased toward business, trade, commodities, and policy news. This domain bias affects evaluation because distributional similarity is learned from local context rather than from dictionary meaning.")

    add_heading(doc, "2. Model Training")
    add_table(doc, ["Model", "Objective", "Vocab", "Dim.", "Window", "Epochs", "Final loss"], [
        ["CBOW", "context -> center", "31,081", "64", "2", "10", "6462.34"],
        ["Skip-Gram", "center -> context", "31,081", "64", "2", "10", "37952.81"],
    ], [0.95, 1.55, 0.8, 0.65, 0.75, 0.75, 0.9])
    add_paragraph(doc, "The loss values are not directly comparable because the two models generate different training targets. CBOW has one target for each center position, while Skip-Gram creates many center-context training pairs. Therefore, the three downstream evaluation tasks are used as the main comparison.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "3. Evaluation Methodology")
    add_paragraph(doc, "KNN evaluation computes cosine similarity between each query vector and all vocabulary vectors, excludes the query itself, and records the top 10 nearest words. The average of those 10 similarities summarizes each query, and the mean across all covered queries summarizes the model.")
    add_paragraph(doc, "SimLex-999 evaluation compares embedding similarity with human similarity scores. Cosine similarity is mapped from [-1, 1] to [0, 10] using scaled = (cosine + 1) * 5. Spearman correlation is then computed between human scores and model scores.")
    add_paragraph(doc, "Analogical reasoning evaluates vector arithmetic. For A:B :: C:D, the predicted vector is embedding(B) - embedding(A) + embedding(C). The closest vocabulary vector is selected by cosine similarity while excluding A, B, and C.")

    add_heading(doc, "4. K-Nearest Neighbor Evaluation")
    add_paragraph(doc, f"Both models cover 49 of the 50 required query words. CBOW obtains an overall average top-10 cosine similarity of {cbow['knn_avg']:.6f}; Skip-Gram obtains {skip['knn_avg']:.6f}. The four detailed words selected by student ID are {', '.join(detail_words)}.")
    add_picture(doc, ASSETS / "knn_overall.png", "Figure 1. Overall average KNN similarity.")
    add_picture(doc, ASSETS / "knn_detail_words.png", "Figure 2. Average KNN similarity for the four selected query words.")

    rows = [[word, neighbor_list(cbow["knn_detail"], word), neighbor_list(skip["knn_detail"], word)] for word in detail_words]
    add_table(doc, ["Query", "CBOW top-10 neighbors", "Skip-Gram top-10 neighbors"], rows, [0.75, 3.05, 3.05])
    add_paragraph(doc, "The neighbors show contextual relatedness more than strict synonymy. For company, Skip-Gram retrieves companies, while CBOW retrieves firm among several noisy news tokens. For help, both models find words related to assistance or intervention. For student, both models are weaker, likely because Reuters is not education-centered.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "5. SimLex-999 Golden Standard Evaluation")
    add_paragraph(doc, f"Both models cover 680 of 999 SimLex pairs. CBOW obtains Spearman correlation {cbow['simlex_spearman']:.6f}; Skip-Gram obtains {skip['simlex_spearman']:.6f}. The low correlations indicate weak alignment with human lexical similarity judgments.")
    add_table(doc, ["Model", "Covered pairs", "Spearman", "Interpretation"], [
        ["CBOW", f"{cbow['simlex_covered']} / 999", f"{cbow['simlex_spearman']:.6f}", "weak positive alignment"],
        ["Skip-Gram", f"{skip['simlex_covered']} / 999", f"{skip['simlex_spearman']:.6f}", "slightly better weak alignment"],
    ], [1.0, 1.2, 1.1, 2.3])
    add_picture(doc, ASSETS / "simlex_sample_scatter.png", "Figure 3. SimLex sampled pairs: standard score vs. scaled embedding score.")

    sim_rows = []
    for c, s in zip(cbow["simlex_sample"].head(20).itertuples(), skip["simlex_sample"].head(20).itertuples()):
        sim_rows.append([f"{c.word1}-{c.word2}", f"{float(c.standard_similarity):.2f}", f"{float(c.scaled_similarity):.2f}", f"{float(s.scaled_similarity):.2f}"])
    add_table(doc, ["Pair", "Standard", "CBOW scaled", "Skip scaled"], sim_rows, [1.65, 0.85, 1.05, 1.05])
    add_paragraph(doc, "The sample shows that many model scores cluster near the middle of the [0, 10] range. Some high-similarity pairs, such as essential-necessary and get-buy, receive reasonably high model scores. But weakly related pairs such as water-salt and boat-car are overestimated because Reuters topical context can make words co-occur without making them true semantic synonyms.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. Analogical Reasoning Evaluation")
    add_paragraph(doc, f"Analogy reasoning is the strictest evaluation. Both models cover 8,591 analogy questions. CBOW answers {cbow['analogy_correct']} correctly, accuracy {cbow['analogy_accuracy']:.6f}. Skip-Gram answers {skip['analogy_correct']} correctly, accuracy {skip['analogy_accuracy']:.6f}.")
    add_table(doc, ["Model", "Covered", "Correct", "Accuracy"], [
        ["CBOW", str(cbow["analogy_covered"]), str(cbow["analogy_correct"]), f"{cbow['analogy_accuracy']:.6f}"],
        ["Skip-Gram", str(skip["analogy_covered"]), str(skip["analogy_correct"]), f"{skip['analogy_accuracy']:.6f}"],
    ], [1.0, 1.0, 1.0, 1.0])
    add_picture(doc, ASSETS / "analogy_category_accuracy.png", "Figure 4. Analogy accuracy by category.")

    cat_rows = []
    for _, row in cbow["analogy_category"].iterrows():
        srow = skip["analogy_category"][skip["analogy_category"]["category"] == row["category"]].iloc[0]
        cat_rows.append([row["category"], f"{int(row['correct'])}/{int(row['covered_questions'])}", f"{float(row['accuracy']):.4f}", f"{int(srow['correct'])}/{int(srow['covered_questions'])}", f"{float(srow['accuracy']):.4f}"])
    add_table(doc, ["Category", "CBOW", "CBOW acc.", "Skip", "Skip acc."], cat_rows, [2.2, 1.0, 0.9, 1.0, 0.9])
    add_paragraph(doc, "Most analogy categories are near zero. CBOW performs best on nationality adjectives, while Skip-Gram distributes a few correct answers across more categories. This indicates that the vector space captures local similarity but not stable relation directions.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Sample Analogy Predictions")
    ana_rows = []
    for c, s in zip(cbow["analogy_sample"].head(10).itertuples(), skip["analogy_sample"].head(10).itertuples()):
        ana_rows.append([c.category, f"{c.word_a}:{c.word_b} :: {c.word_c}:?", c.expected, c.predicted, s.predicted])
    add_table(doc, ["Category", "Analogy", "Expected", "CBOW", "Skip-Gram"], ana_rows, [1.55, 2.25, 1.0, 1.0, 1.0])
    add_paragraph(doc, "The sampled predictions are mostly incorrect and often return Reuters-specific tokens. This is a stronger failure than SimLex because analogical reasoning requires the difference between two word vectors to be meaningful and reusable across other word pairs.")

    add_heading(doc, "Vector Plot Examples", 2)
    add_picture(doc, RESULTS / "cbow_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png", "Figure 5. CBOW predicted vector and nearest searched vector.")
    add_picture(doc, RESULTS / "skipgram_analogy_plots" / "capital-common-countries_athens_greece_baghdad.png", "Figure 6. Skip-Gram predicted vector and nearest searched vector.")
    add_paragraph(doc, "The full vector plot PNGs and CSV files for five categories per model are included in the results folder. The plotted lines compare the predicted vector with the nearest searched word vector across 64 dimensions.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. Overall Comparison and Discussion")
    add_paragraph(doc, "Skip-Gram is consistently but slightly stronger than CBOW in this run. It has the higher KNN average, the higher SimLex Spearman correlation, and the higher analogy accuracy. However, the margins are small, so the practical conclusion is not that Skip-Gram solves the tasks, but that it performs a little better under this setup.")
    add_paragraph(doc, "The pattern across tasks is important. KNN is the most favorable evaluation because it only requires local neighborhoods to be plausible. SimLex is harder because it asks whether the model ranking agrees with human semantic similarity. Analogy is hardest because it asks whether semantic and grammatical relationships become stable vector offsets.")
    add_paragraph(doc, "The main limitations are the Reuters corpus domain, the small embedding dimension, and the plain Word2Vec objective. Larger and more balanced training data, negative sampling, rare-token filtering, more epochs, and hyperparameter tuning would likely improve the results.")
    add_heading(doc, "8. Conclusion")
    add_paragraph(doc, "This project successfully trains and evaluates CBOW and Skip-Gram embeddings. The KNN results demonstrate that both models learn useful distributional neighborhoods. The SimLex and analogy results show that the embeddings do not strongly match human semantic similarity or exact relation arithmetic. Skip-Gram is the better model in this experiment, but both models remain limited by data scale and training setup.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    data = load_results()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    (REPORT_DIR / "project1_2_report.md").write_text(build_markdown(data), encoding="utf-8")
    build_docx(data, REPORT_DIR / "project1_2_report.docx")
    print(REPORT_DIR / "project1_2_report.md")
    print(REPORT_DIR / "project1_2_report.docx")


if __name__ == "__main__":
    main()
